# -*- coding: utf-8 -*-
"""
提案書 (Google Docs 取り込み用 .docx)
「生成AI実務活用 5日間ハンズオン研修」導入提案書
"""

from docx import Document
from docx.shared import Pt, Cm, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# --- color palette (matching the brand) ---
PRIMARY      = RGBColor(0x6C, 0x47, 0xFF)
PRIMARY_DARK = RGBColor(0x4A, 0x2B, 0xD4)
ACCENT       = RGBColor(0xFF, 0x6B, 0x35)
TEXT         = RGBColor(0x1A, 0x1A, 0x2E)
MUTED        = RGBColor(0x6B, 0x72, 0x80)
SUCCESS      = RGBColor(0x16, 0xA3, 0x4A)
HEADER_BG    = "F8F7FC"
HIGHLIGHT_BG = "FFF4EE"
SUBTLE_BG    = "F6F5FA"

JP_FONT_HEAD = "Yu Gothic"
JP_FONT_BODY = "Yu Mincho"  # use mincho for body to feel like a formal proposal
JP_FONT_SANS = "Yu Gothic"  # for tables / labels


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def set_run_font(run, font_name=JP_FONT_BODY, size=10.5, color=TEXT, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    # ensure CJK font (Word stores east-asian font separately)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_paragraph(doc, text="", *, font=JP_FONT_BODY, size=10.5, color=TEXT,
                  bold=False, align="left", space_before=0, space_after=6,
                  line_spacing=1.6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        set_run_font(run, font_name=font, size=size, color=color, bold=bold)
    return p


def add_runs(doc, parts, *, align="left", space_before=0, space_after=6,
             line_spacing=1.6, default_font=JP_FONT_BODY, default_size=10.5,
             default_color=TEXT):
    """parts: list of (text, opts_dict)"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    for text, opts in parts:
        opts = opts or {}
        run = p.add_run(text)
        set_run_font(
            run,
            font_name=opts.get("font", default_font),
            size=opts.get("size", default_size),
            color=opts.get("color", default_color),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
        )
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="DDDDDD", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        tcBorders.append(b)


def set_cell_text(cell, text, *, font=JP_FONT_SANS, size=10, color=TEXT,
                  bold=False, align="left", line_spacing=1.4):
    cell.text = ""  # clear default empty paragraph text
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    set_run_font(run, font_name=font, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_paragraph_left_border(p, color="6C47FF", size=24, space=4):
    """Apply a left bar border to a paragraph (used for h2-style headings)."""
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def set_paragraph_shading(p, hex_color):
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_h1(doc, text):
    p = add_paragraph(doc, text, font=JP_FONT_HEAD, size=18, color=PRIMARY_DARK,
                      bold=True, space_before=18, space_after=10, line_spacing=1.4)
    pf = p.paragraph_format
    pf.keep_with_next = True
    return p


def add_h2(doc, num, text):
    """Section heading like '1. 提案概要' with a left-bar accent."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(20)
    pf.space_after = Pt(10)
    pf.line_spacing = 1.35
    pf.keep_with_next = True
    # number in primary color
    r1 = p.add_run(f"{num}.  ")
    set_run_font(r1, font_name=JP_FONT_HEAD, size=15, color=PRIMARY, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, font_name=JP_FONT_HEAD, size=15, color=PRIMARY_DARK, bold=True)
    set_paragraph_left_border(p, color="6C47FF", size=24, space=8)
    return p


def add_h3(doc, text):
    p = add_paragraph(doc, text, font=JP_FONT_HEAD, size=12, color=PRIMARY_DARK,
                      bold=True, space_before=12, space_after=6, line_spacing=1.4)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    """Markdown-style bullet using the Word built-in 'List Bullet' style."""
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.55
    # support **bold** inline
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = p.add_run(part)
        set_run_font(run, font_name=JP_FONT_BODY, size=10.5, color=TEXT,
                     bold=(i % 2 == 1))
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.55
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = p.add_run(part)
        set_run_font(run, font_name=JP_FONT_BODY, size=10.5, color=TEXT,
                     bold=(i % 2 == 1))
    return p


def add_table(doc, rows, *, header=True, col_widths=None, header_bg=HEADER_BG,
              col_aligns=None):
    """rows: list[list[str]]; first row treated as header if header=True."""
    n_rows = len(rows)
    n_cols = len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    # set widths
    if col_widths is None:
        col_widths = [Cm(16 / n_cols)] * n_cols
    for j, w in enumerate(col_widths):
        for i in range(n_rows):
            table.cell(i, j).width = w

    aligns = col_aligns or ["left"] * n_cols

    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.cell(i, j)
            is_header = header and i == 0
            set_cell_borders(cell, color="DDDDDD", size=4)
            if is_header:
                shade_cell(cell, header_bg)
                set_cell_text(cell, txt,
                              font=JP_FONT_HEAD, size=10.5,
                              color=PRIMARY_DARK, bold=True,
                              align=aligns[j])
            else:
                # support inline **bold**
                cell.text = ""
                p = cell.paragraphs[0]
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.4
                if aligns[j] == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif aligns[j] == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                parts = txt.split("**")
                for k, part in enumerate(parts):
                    if part == "":
                        continue
                    run = p.add_run(part)
                    set_run_font(run, font_name=JP_FONT_SANS, size=10,
                                 color=TEXT, bold=(k % 2 == 1))
    return table


def add_callout(doc, lines, *, bg=HIGHLIGHT_BG, accent_color="FF6B35"):
    """A boxed callout block — single-cell table with shading + left bar."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.0)
    set_cell_borders(cell, color="EEEEEE", size=4)
    shade_cell(cell, bg)
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # set padding
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", 160), ("bottom", 160), ("left", 200), ("right", 200)):
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    for idx, (label, text) in enumerate(lines):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0 if idx == 0 else 4)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.55
        if label:
            run = p.add_run(label)
            set_run_font(run, font_name=JP_FONT_HEAD, size=10.5,
                         color=PRIMARY_DARK, bold=True)
        if text:
            run = p.add_run(text)
            set_run_font(run, font_name=JP_FONT_BODY, size=10.5, color=TEXT)

    # left bar via paragraph border on first paragraph (workaround: use cell border-left)
    tcPr2 = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr2.find(qn("w:tcBorders"))
    if tcBorders is not None:
        left = tcBorders.find(qn("w:left"))
        if left is not None:
            left.set(qn("w:sz"), "24")
            left.set(qn("w:color"), accent_color)

    # spacing after the callout
    add_paragraph(doc, "", space_before=0, space_after=6)
    return table


# ---------------------------------------------------------------------------
# build document
# ---------------------------------------------------------------------------

doc = Document()

# A4 with 25mm margins
section = doc.sections[0]
section.page_width  = Mm(210)
section.page_height = Mm(297)
section.top_margin    = Mm(22)
section.bottom_margin = Mm(22)
section.left_margin   = Mm(22)
section.right_margin  = Mm(22)

# default style
style = doc.styles["Normal"]
style.font.name = JP_FONT_BODY
style.font.size = Pt(10.5)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.insert(0, rFonts)
rFonts.set(qn("w:eastAsia"), JP_FONT_BODY)
rFonts.set(qn("w:ascii"), JP_FONT_BODY)
rFonts.set(qn("w:hAnsi"), JP_FONT_BODY)


# =============================================================================
# COVER
# =============================================================================
add_paragraph(doc, "提  案  書",
              font=JP_FONT_HEAD, size=12, color=MUTED, bold=True,
              align="center", space_before=60, space_after=18, line_spacing=1.4)

add_paragraph(doc, "「生成AI実務活用 5日間ハンズオン研修」",
              font=JP_FONT_HEAD, size=22, color=PRIMARY_DARK, bold=True,
              align="center", space_before=0, space_after=8, line_spacing=1.5)
add_paragraph(doc, "導入提案",
              font=JP_FONT_HEAD, size=22, color=PRIMARY_DARK, bold=True,
              align="center", space_before=0, space_after=14, line_spacing=1.5)

add_paragraph(doc, "Claude Chat を活用した職員向け業務効率化プログラム",
              font=JP_FONT_BODY, size=12, color=MUTED,
              align="center", space_before=0, space_after=42, line_spacing=1.4)

# meta box
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_table.autofit = False
meta_rows = [
    ("提  案  者", "（提案者名・組織名を記入）"),
    ("提  出  日", "2026年5月"),
    ("提  出  先", "（提出先・部署を記入）"),
    ("教 材 U R L", "https://atsunaricoda-maker.github.io/claude-chat-course/"),
]
for i, (label, val) in enumerate(meta_rows):
    c1 = meta_table.cell(i, 0)
    c2 = meta_table.cell(i, 1)
    c1.width = Cm(4.5)
    c2.width = Cm(11.5)
    set_cell_borders(c1, color="EEEEEE", size=4)
    set_cell_borders(c2, color="EEEEEE", size=4)
    shade_cell(c1, HEADER_BG)
    set_cell_text(c1, label, font=JP_FONT_HEAD, size=10.5,
                  color=PRIMARY_DARK, bold=True, align="left")
    set_cell_text(c2, val, font=JP_FONT_BODY, size=10.5, color=TEXT, align="left")

# cover bottom note
add_paragraph(doc, "本提案書および別添の配布用スライド資料は、Web教材「Claude Chat 5日間ハンズオン講座」"
                    "を貴組織の研修プログラムとして導入いただくためのご提案資料です。",
              font=JP_FONT_BODY, size=10, color=MUTED,
              align="center", space_before=42, space_after=0, line_spacing=1.7)

# page break to TOC
p_break = doc.add_paragraph()
p_break.add_run().add_break(WD_BREAK.PAGE)
p_break_after_cover = True  # marker

# =============================================================================
# TABLE OF CONTENTS
# =============================================================================
add_h1(doc, "目  次")
toc_items = [
    ("1.",  "提案概要"),
    ("2.",  "背景・課題認識"),
    ("3.",  "本研修の特徴"),
    ("4.",  "カリキュラム詳細"),
    ("5.",  "期待される成果（受講者個人）"),
    ("6.",  "期待される成果（組織・社会）"),
    ("7.",  "投資対効果（ROI 試算）"),
    ("8.",  "実施体制・必要要件"),
    ("9.",  "リスクと対策"),
    ("10.", "実施スケジュール提案"),
    ("11.", "結語"),
]
for num, title in toc_items:
    add_runs(doc, [
        (f"  {num}  ", {"font": JP_FONT_HEAD, "color": PRIMARY, "bold": True}),
        (title, {"font": JP_FONT_HEAD, "size": 11, "color": TEXT}),
    ], space_before=0, space_after=4, line_spacing=1.6)

p_break = doc.add_paragraph()
p_break.add_run().add_break(WD_BREAK.PAGE)


# =============================================================================
# SECTION 1
# =============================================================================
add_h2(doc, "1", "提案概要")
add_paragraph(doc,
    "本研修は、生成AI「Claude Chat」を職員の日常業務に組み込むための、ハンズオン中心の体系的研修プログラムです。"
    "従来の「単発の体験会」「講義中心の研修」では業務定着に至らないという課題を、5日間の段階的設計と職種別教材で解決します。",
    align="justify")

add_table(doc,
    [
        ["項目", "内容"],
        ["名称", "Claude Chat 5日間ハンズオン研修"],
        ["対象", "職種を問わない全職員（非エンジニア中心）"],
        ["期間", "全5日間 ／ 各2時間 ／ 合計10時間"],
        ["形式", "ハンズオン演習主体（オンライン or 対面）"],
        ["想定受講者数", "1回あたり10〜30名"],
        ["教材", "Web上で全公開済み・常時アクセス可能"],
        ["研修後フォロー", "30日定着ロードマップ付き"],
    ],
    col_widths=[Cm(4.5), Cm(11.5)],
    col_aligns=["center", "left"])


# =============================================================================
# SECTION 2
# =============================================================================
add_h2(doc, "2", "背景・課題認識")

add_h3(doc, "社会的背景")
add_bullet(doc, "2026年時点で生成AI（Claude等）の業務活用は世界的に標準化")
add_bullet(doc, "OECD加盟国の行政機関で生成AI活用率は前年比で大幅に増加")
add_bullet(doc, "日本の行政機関でも導入検討が急増しているが、現場活用には大きなギャップ")

add_h3(doc, "現場が抱える3つの壁")
add_numbered(doc, "**「研修受けたけど結局使えていない」**— 一般的な生成AI研修受講者の半数以上が業務に活かせていないという調査結果")
add_numbered(doc, "**ITリテラシーの個人差**— ITに不慣れな職員ほどハードルが高く、組織全体での推進が困難")
add_numbered(doc, "**職種別ニーズへの対応不足**— 部署・職種ごとに使い方が異なるが、画一的な研修になりがち")

add_callout(doc, [
    ("本提案の意義：", "ハンズオン中心の設計と職種別教材により、「使える状態」まで到達した受講者を確実に生み出します。"
                       "30日のフォローロードマップで業務定着を保証し、研修投資が確実に成果につながる構造を採用しています。"),
])


# =============================================================================
# SECTION 3
# =============================================================================
add_h2(doc, "3", "本研修の特徴")

add_paragraph(doc, "一般的な生成AI研修と本研修との比較を以下に整理します。",
              align="justify", space_after=8)

add_table(doc,
    [
        ["観点",        "一般的な研修",        "本研修"],
        ["形式",        "講義中心",            "ハンズオン90%"],
        ["プロンプト例", "一律のサンプル",      "営業・マーケ・人事・経理・総務・管理職の6職種別"],
        ["期間構成",    "1日完結",             "5日間で段階的に習得"],
        ["研修後",      "フォローなし",        "30日ロードマップ付き"],
        ["教材形態",    "紙資料",              "Web常時アクセス可能・スマホ対応"],
        ["再現性",      "講師の力量に依存",    "コピペ可能なプロンプト集で誰でも再実施可能"],
        ["UI言語対応",  "英語表記のみ",        "英語UIラベルに日本語併記"],
    ],
    col_widths=[Cm(3.5), Cm(5.0), Cm(7.5)],
    col_aligns=["center", "left", "left"])


# =============================================================================
# SECTION 4
# =============================================================================
add_h2(doc, "4", "カリキュラム詳細")

# Day 1
add_h3(doc, "Day 1：Claude を自分仕様にする")
add_paragraph(doc,
    "Project機能・Memory機能を使い、Claudeに自分の業務文脈・文体を学習させ、"
    "「毎回説明しなくていい相棒」を構築します。", align="justify")
add_bullet(doc, "アカウント・UI 案内（15分）")
add_bullet(doc, "Project 作成と指示文設定（30分）")
add_bullet(doc, "過去の文書から文体を学習（35分）")
add_bullet(doc, "Memory で個人情報を記憶（25分）")
add_bullet(doc, "振り返り・宿題（15分）")

# Day 2
add_h3(doc, "Day 2：「読む」「まとめる」を委任する")
add_paragraph(doc,
    "PDF・議事録・録音を Claude に処理させ、資料処理時間を 1/10 に短縮します。", align="justify")
add_bullet(doc, "PDF を「自分の視点で」要約")
add_bullet(doc, "議事録 → ToDo リスト・共有メール変換")
add_bullet(doc, "Artifacts で 1枚A4サマリー作成")
add_bullet(doc, "会話で資料を編集")

# Day 3
add_h3(doc, "Day 3：データとスライドを自動化する")
add_paragraph(doc,
    "Skill 機能（/xlsx、/pptx）で Excel職人・PPT職人を卒業します。", align="justify")
add_bullet(doc, "Skill の概念・カタログ・ベネフィット理解")
add_bullet(doc, "Skill の設定方法・FAQ")
add_bullet(doc, "/xlsx で月次レポート自動生成")
add_bullet(doc, "/pptx で15分プレゼン自動作成")
add_bullet(doc, "ブランド準拠デザイン")

# Day 4
add_h3(doc, "Day 4：周辺ツールと繋いで「アプリ往復」を消す")
add_paragraph(doc,
    "Connector 機能で Gmail / Calendar / Drive を統合し、横断的な情報処理を実現します。", align="justify")
add_bullet(doc, "Connector の概念・カタログ・ベネフィット理解")
add_bullet(doc, "接続方法・FAQ・プライバシー運用")
add_bullet(doc, "Gmail 要約・返信ドラフト")
add_bullet(doc, "Calendar 週次プランニング")
add_bullet(doc, "Drive 横断検索")
add_bullet(doc, "3ツール統合ブリーフィング")

# Day 5
add_h3(doc, "Day 5：自分の業務に組み込む")
add_paragraph(doc,
    "学んだ機能を自分の定常業務に組み込む設計を、各受講者が自ら作成します。", align="justify")
add_bullet(doc, "業務棚卸しワーク（A〜D 分類）")
add_bullet(doc, "1業務をフル Claude 化する詳細設計")
add_bullet(doc, "同僚向け 3分デモシナリオ作成")
add_bullet(doc, "30日定着ロードマップ作成")


# =============================================================================
# SECTION 5
# =============================================================================
add_h2(doc, "5", "期待される成果（受講者個人）")

add_h3(doc, "定量的な業務時間削減（試算）")
add_table(doc,
    [
        ["業務",              "現状",                 "Claude化後",           "削減時間/月"],
        ["議事録作成",         "90分 × 4回 = 360分",   "10分 × 4回 = 40分",    "320分"],
        ["月次レポート作成",    "240分 × 1回",         "60分 × 1回",           "180分"],
        ["メール返信",         "30分 × 20日 = 600分",  "10分 × 20日 = 200分",  "400分"],
        ["提案書作成",         "180分 × 4本 = 720分",  "60分 × 4本 = 240分",   "480分"],
        ["議事録共有メール",    "30分 × 4回",          "5分 × 4回",            "100分"],
        ["**合計（月）**",      "**約 33時間**",        "**約 12時間**",        "**約 21時間**"],
    ],
    col_widths=[Cm(3.5), Cm(4.5), Cm(4.5), Cm(3.5)],
    col_aligns=["left", "left", "left", "right"])

add_h3(doc, "定性的な成果")
add_bullet(doc, "**文書品質の標準化** — ベテランと若手の差が縮まる、属人化の解消")
add_bullet(doc, "**「白紙が怖い」状態からの脱却** — 初稿の負担が減り、創造的業務に時間を使える")
add_bullet(doc, "**自走力の獲得** — 自分でカスタマイズできる、ツール変化に対応できる")
add_bullet(doc, "**社内講師化** — 同僚に説明・伝授でき、研修が組織内に伝播する")


# =============================================================================
# SECTION 6
# =============================================================================
add_h2(doc, "6", "期待される成果（組織・社会）")

add_h3(doc, "組織への波及効果")
add_bullet(doc, "1人あたり月21時間 × 受講者数 = 大きな時間創出")
add_bullet(doc, "創出された時間を市民サービス向上・企画立案・新規事業に転用可能")
add_bullet(doc, "部署横断のベストプラクティス共有による組織学習")
add_bullet(doc, "採用・新人育成期間の短縮")
add_bullet(doc, "残業時間削減・ワークライフバランス改善")

add_h3(doc, "社会的便益")
add_bullet(doc, "職員の働き方改革推進")
add_bullet(doc, "公務効率化による財政負担の軽減")
add_bullet(doc, "生成AI活用の先進事例として他自治体・他組織への波及")
add_bullet(doc, "DX推進指標（KPI）への直接的貢献")
add_bullet(doc, "市民への迅速な対応・サービス品質向上")


# =============================================================================
# SECTION 7
# =============================================================================
add_h2(doc, "7", "投資対効果（ROI 試算）")
add_paragraph(doc,
    "受講者30名規模の場合の試算です。実際の規模・条件により調整可能です。",
    align="justify")

add_table(doc,
    [
        ["項目",              "金額（年間）",     "内訳"],
        ["投資",              "約 100 万円",      "講師費 + ツール費 + 機会費用（教材費は ¥0）"],
        ["**創出付加価値**",   "**約 1,440 万円**", "受講者 1人 × 192時間 × 2,500円換算 × 30名"],
    ],
    col_widths=[Cm(3.5), Cm(4.5), Cm(8.0)],
    col_aligns=["left", "right", "left"])

add_h3(doc, "計算根拠")
add_bullet(doc, "1人あたりの月間時間削減：21時間（前項試算）")
add_bullet(doc, "年間時間削減：21時間 × 12ヶ月 = **252時間/人/年**（保守的に192時間で計算）")
add_bullet(doc, "時給換算：年収500万円相当として時給 2,500円")
add_bullet(doc, "1人あたり付加価値：192時間 × 2,500円 = **480,000円/年**")
add_bullet(doc, "30名：480,000円 × 30 = **14,400,000円/年**")

add_callout(doc, [
    ("回収期間：約1ヶ月。", "投資100万円に対し、月間120万円の付加価値創出により、初月でほぼ回収完了する高ROI施策です。"),
])


# =============================================================================
# SECTION 8
# =============================================================================
add_h2(doc, "8", "実施体制・必要要件")

add_h3(doc, "必要なもの")
add_bullet(doc, "受講者用 PC（インターネット接続）")
add_bullet(doc, "Claude.ai アカウント（無料プラン〜Enterprise）")
add_bullet(doc, "集合形式の場合：会議室・プロジェクター")
add_bullet(doc, "オンライン形式の場合：ビデオ会議ツール（Zoom／Teams等）")

add_h3(doc, "事前準備（推奨）")
add_bullet(doc, "受講者の業務カテゴリ事前ヒアリング（職種別プロンプトの選定に活用）")
add_bullet(doc, "機密情報取扱いポリシーのすり合わせ")
add_bullet(doc, "情シス部門との接続確認（Connector 利用時）")
add_bullet(doc, "実業務の素材（メール・議事録・データ等）の事前準備依頼")

add_h3(doc, "講師")
add_bullet(doc, "本カリキュラム作成者または認定講師")
add_bullet(doc, "非エンジニア研修の経験者")
add_bullet(doc, "受講者の質問にリアルタイムで回答できる体制")


# =============================================================================
# SECTION 9
# =============================================================================
add_h2(doc, "9", "リスクと対策")

add_table(doc,
    [
        ["リスク",                "対策"],
        ["機密情報の取扱い",
         "個人情報・機密情報を含む資料は事前ガイドラインで明確化。研修中はダミーデータ・匿名化資料のみ使用。"
         "Day 4 でプライバシー運用ルールを必ず確認。"],
        ["ツール側の仕様変更",
         "Web教材として常時更新可能な構造を採用。月次でメンテナンス実施。UI変更時は即座に教材に反映。"],
        ["受講者間のITリテラシー差",
         "6職種別プロンプトと「困ったら使えるお守りプロンプト集」（Day5付録）で底上げ。クリック箇所を全て明記。"],
        ["業務適用への抵抗",
         "Day5 の業務棚卸しワークで個別最適化。「同僚向けデモ」シナリオ作成で組織内推進者を育成。"],
        ["研修効果の測定",
         "30日後にフォローアップアンケート実施。時間削減効果を実測値で報告。"],
    ],
    col_widths=[Cm(4.5), Cm(11.5)],
    col_aligns=["left", "left"])


# =============================================================================
# SECTION 10
# =============================================================================
add_h2(doc, "10", "実施スケジュール提案")

add_h3(doc, "標準スケジュール（2週間プログラム）")
add_table(doc,
    [
        ["週",       "曜日",  "内容",                                "時間"],
        ["第 1 週", "月",    "Day 1：Claude を自分仕様にする",         "2 時間"],
        ["第 1 週", "水",    "Day 2：「読む」「まとめる」を委任する",   "2 時間"],
        ["第 1 週", "金",    "Day 3：データとスライドを自動化する",      "2 時間"],
        ["第 2 週", "火",    "Day 4：周辺ツールと繋ぐ",                "2 時間"],
        ["第 2 週", "木",    "Day 5：自分の業務に組み込む",             "2 時間"],
        ["実 行 期", "—",   "受講者各自で 30日ロードマップを実行",     "30 日間"],
        ["30日後",   "—",    "フォローアップ（希望者対象）",            "30 分"],
    ],
    col_widths=[Cm(2.5), Cm(2.0), Cm(8.5), Cm(3.0)],
    col_aligns=["center", "center", "left", "right"])

add_h3(doc, "柔軟な調整")
add_bullet(doc, "1日に 2 Day 連続実施で 1週間圧縮も可能")
add_bullet(doc, "1ヶ月に 1 Day ずつのゆったり進行も可能")
add_bullet(doc, "部署別開催・全社一括開催どちらも対応")


# =============================================================================
# SECTION 11
# =============================================================================
add_h2(doc, "11", "結語")

add_callout(doc, [
    ("",
     "本研修は「研修受けたが業務に活かせていない」という典型的な失敗を、ハンズオン中心の設計と"
     "職種別教材で乗り越える設計となっております。"),
    ("",
     "5日間 × 2時間という現実的な時間投資で、職員が自走できる状態に到達し、"
     "その後30日で業務定着を実現する仕組みです。"),
    ("",
     "DX推進・働き方改革・業務効率化の3つの政策目標に同時に貢献する施策として、"
     "ぜひご検討くださいますようお願い申し上げます。"),
], bg="F0EEFC", accent_color="6C47FF")


# Footer / contact
add_paragraph(doc, "", space_before=12, space_after=0)
add_paragraph(doc, "─────────────────────────────",
              font=JP_FONT_BODY, size=10, color=MUTED, align="center",
              space_before=0, space_after=4)
add_paragraph(doc,
    "本提案に関するお問合せは（提案者連絡先）まで",
    font=JP_FONT_BODY, size=10, color=MUTED, align="center",
    space_before=0, space_after=2)
add_paragraph(doc,
    "教材URL：https://atsunaricoda-maker.github.io/claude-chat-course/",
    font=JP_FONT_BODY, size=10, color=PRIMARY_DARK, align="center",
    space_before=0, space_after=0)


out_path = "Claude_Chat_5日間研修_導入提案書.docx"
doc.save(out_path)
print(f"saved: {out_path}")
